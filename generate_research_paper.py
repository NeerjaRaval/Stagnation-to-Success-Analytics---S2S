from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(31, 73, 125)):
    para = doc.add_heading(level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    elif level == 3:
        run.font.size = Pt(11.5)
    return para

def add_paragraph(doc, text, bold_prefix=None, indent=False):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Inches(0.3)
    if bold_prefix:
        run = para.add_run(bold_prefix + " ")
        set_font(run, bold=True, size=11)
    run = para.add_run(text)
    set_font(run, size=11)
    para.paragraph_format.space_after = Pt(6)
    return para

def add_bullet(doc, text, bold_prefix=None, level=1):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.3 * level)
    if bold_prefix:
        r = para.add_run(bold_prefix + ": ")
        set_font(r, bold=True, size=11)
    r = para.add_run(text)
    set_font(r, size=11)
    para.paragraph_format.space_after = Pt(3)

def add_divider(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(4)

def shade_cell(cell, fill_hex="1F497D"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), 'B8CCE4')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, "1F497D")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, bold=True, size=10, color=(255, 255, 255))

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "D9E1F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            shade_cell(cell, fill)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            set_font(run, size=10)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table

# ─── Build Document ──────────────────────────────────────────────────────────

doc = Document()

# ── Page Margins ──
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ══════════════════════════════════════════════════════════════════════════════
# COVER / TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_para.add_run("Stagnation-to-Success Analytics (S2S)")
set_font(tr, name="Calibri", size=22, bold=True, color=(31, 73, 125))

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle_para.add_run(
    "Career Progression and Promotion Gap Analysis\n"
    "for Retention Optimization at Palo Alto Networks"
)
set_font(sr, name="Calibri", size=15, bold=False, italic=True, color=(68, 114, 196))

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

meta_lines = [
    ("Prepared by:",   "Neerja"),
    ("Organization:",  "Palo Alto Networks — HR Analytics Division"),
    ("Date:",          datetime.date.today().strftime("%B %d, %Y")),
    ("Document Type:", "Internship Research Paper"),
    ("Version:",       "1.0 (Final)"),
]
for label, value in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lb = p.add_run(label + "  ")
    set_font(lb, bold=True, size=11, color=(31, 73, 125))
    vr = p.add_run(value)
    set_font(vr, size=11)

doc.add_paragraph()
add_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Abstract", level=1)
abstract_text = (
    "Employee attrition is one of the most costly and disruptive challenges facing modern "
    "technology organizations. Traditional predictive attrition models focus on identifying "
    "who is likely to leave, but rarely explain the structural career reasons driving that "
    "disengagement. This research paper presents the Stagnation-to-Success (S2S) Analytics "
    "framework — a data-driven approach developed during an internship at Palo Alto Networks — "
    "that shifts the paradigm from reactive attrition prediction to proactive career-centric "
    "retention management.\n\n"
    "Using a rich employee dataset comprising 30 career and satisfaction attributes, S2S applies "
    "custom feature engineering to construct four novel career intelligence metrics: the "
    "Promotion Gap Ratio, Role Stagnation Index, Training Intensity Score, and Manager "
    "Stagnation Index. These metrics power an unsupervised K-Means clustering pipeline that "
    "segments employees into four distinct career trajectory archetypes: Fast-Track Performers, "
    "Stable Long-Term Contributors, Early-Career Explorers, and Promotion-Stalled High-Risk "
    "Profiles. A supervised Random Forest Classifier further quantifies stagnation-driven "
    "attrition risk for each employee profile.\n\n"
    "The resulting analytics platform — delivered as a full-stack web application (React frontend "
    "+ FastAPI backend) — provides HR stakeholders with actionable dashboards for identifying "
    "retention opportunities before disengagement occurs, enabling Palo Alto Networks to move "
    "toward proactive, personalized workforce management."
)
add_paragraph(doc, abstract_text)
doc.add_paragraph()

kw_para = doc.add_paragraph()
kw_run = kw_para.add_run("Keywords: ")
set_font(kw_run, bold=True, size=11)
kw_val = kw_para.add_run(
    "Career Intelligence, Employee Retention, Promotion Gap Analysis, "
    "K-Means Clustering, Random Forest, HR Analytics, Attrition Prediction, Workforce Management"
)
set_font(kw_val, italic=True, size=11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Table of Contents", level=1)
toc_items = [
    ("1.", "Introduction", "3"),
    ("2.", "Problem Statement", "3"),
    ("3.", "Literature Review", "4"),
    ("4.", "Dataset Description", "5"),
    ("5.", "Data Science Methodology", "6"),
    ("  5.1", "Feature Engineering", "6"),
    ("  5.2", "Data Preprocessing", "7"),
    ("  5.3", "Career Path Clustering (Unsupervised Learning)", "7"),
    ("  5.4", "Cluster Interpretation & Labeling", "8"),
    ("  5.5", "Attrition Risk Classification", "8"),
    ("  5.6", "Retention Opportunity Identification", "9"),
    ("6.", "System Architecture & Implementation", "9"),
    ("7.", "Key Performance Indicators (KPIs)", "11"),
    ("8.", "Results and Discussion", "11"),
    ("9.", "Recommendations", "12"),
    ("10.", "Conclusion", "13"),
    ("11.", "References", "13"),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
    nr = p.add_run(f"{num}  {title}")
    set_font(nr, size=11, bold=("." not in num[1:]))
    pr = p.add_run(f"\t{page}")
    set_font(pr, size=11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Introduction", level=1)
add_paragraph(doc,
    "In the highly competitive technology sector, retaining top talent is as strategically "
    "critical as acquiring it. Palo Alto Networks, a global leader in cybersecurity, invests "
    "substantially in recruiting and developing specialized professionals. When skilled employees "
    "leave, the organization bears not only direct replacement costs — estimated at 50–200% of "
    "annual salary — but also the hidden costs of lost institutional knowledge, diminished team "
    "productivity, and disrupted project continuity."
)
add_paragraph(doc,
    "A critical insight, often overlooked by conventional HR analytics, is that many employees "
    "do not leave because of immediate dissatisfaction. Instead, they disengage gradually due to "
    "structural career issues: prolonged gaps between promotions, role stagnation, insufficient "
    "skill development opportunities, and managerial instability. By the time these employees "
    "appear on a predictive attrition radar, the disengagement is already deeply rooted."
)
add_paragraph(doc,
    "This research introduces the Stagnation-to-Success (S2S) Analytics Framework — a "
    "comprehensive data science solution that diagnoses these structural career issues at an "
    "organizational scale, enabling HR teams to intervene proactively before talent is lost."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Problem Statement", level=1)
add_paragraph(doc,
    "Palo Alto Networks currently lacks a systematic, data-driven mechanism to:"
)
add_bullet(doc, "Visualize and understand career progression patterns across departments and roles.")
add_bullet(doc, "Identify employees experiencing promotion stagnation before they exhibit disengagement behavior.")
add_bullet(doc, "Capture early retention opportunities for employees who are still engaged but structurally at risk.")
add_bullet(doc, "Provide HR managers with personalized, actionable career intervention recommendations.")

doc.add_paragraph()
add_paragraph(doc,
    "Without these capabilities, retention strategies remain generic and reactive — applied "
    "broadly rather than tailored to individual career needs. This project directly addresses "
    "this gap by building a career trajectory intelligence system grounded in machine learning "
    "and real employee data."
)
add_paragraph(doc,
    "The central research question this project answers is: "
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
rq = p.add_run(
    '"What structural career patterns — particularly around promotion gaps and role stagnation — '
    'place employees at elevated retention risk, and how can these patterns be identified and '
    'acted upon before disengagement occurs?"'
)
set_font(rq, italic=True, size=11, color=(31, 73, 125))

# ══════════════════════════════════════════════════════════════════════════════
# 3. LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "3. Literature Review", level=1)
add_paragraph(doc,
    "Employee attrition prediction has been an active area of HR analytics research. "
    "Existing literature broadly falls into two categories: predictive attrition modeling "
    "and career progression analytics."
)

add_heading(doc, "3.1 Predictive Attrition Modeling", level=2)
add_paragraph(doc,
    "Numerous studies have applied machine learning algorithms — including logistic regression, "
    "decision trees, support vector machines, and gradient boosting — to predict attrition "
    "probability from employee survey and HR data (Punnoose & Ajit, 2016; Zhao et al., 2018). "
    "The IBM HR Analytics Employee Attrition Dataset, widely used in benchmarking, demonstrates "
    "that features such as OverTime, MonthlyIncome, and JobSatisfaction are among the most "
    "predictive. However, these models primarily answer who will leave, not why."
)

add_heading(doc, "3.2 Career Progression Analytics", level=2)
add_paragraph(doc,
    "Scholarship on internal labor markets (Doeringer & Piore, 1971) and career mobility "
    "(Rosenbaum, 1979) established that promotion velocity and role transitions significantly "
    "influence employee commitment and organizational attachment. More recent studies in "
    "organizational behavior confirm that perceived career stagnation is a major antecedent "
    "of turnover intention (Ng & Feldman, 2014)."
)

add_heading(doc, "3.3 Clustering for HR Segmentation", level=2)
add_paragraph(doc,
    "Unsupervised clustering methods — particularly K-Means — have been applied to HR data "
    "to segment employees into behavioral personas (Saradhi & Palshikar, 2011). Such "
    "segmentation enables targeted retention strategies rather than uniform, one-size-fits-all "
    "interventions. The S2S framework builds on this foundation by combining unsupervised "
    "clustering with novel stagnation-specific feature engineering."
)

add_heading(doc, "3.4 Research Gap", level=2)
add_paragraph(doc,
    "Despite extensive literature on attrition prediction and career analytics, no widely "
    "adopted framework integrates both into a unified, deployable analytics platform with "
    "real-time dashboard capabilities. The S2S project addresses this gap by delivering a "
    "production-ready system that moves from data science insights to actionable HR decisions."
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. DATASET DESCRIPTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "4. Dataset Description", level=1)
add_paragraph(doc,
    "The analysis is based on the Palo Alto Networks Employee Dataset, a structured tabular "
    "dataset containing demographic, compensation, satisfaction, and career history records "
    "for all active and former employees. The dataset is stored as a CSV file "
    "(Palo Alto Networks.csv) containing 1,470 employee records and 30 feature columns."
)
doc.add_paragraph()

dataset_headers = ["Field Name", "Data Type", "Description"]
dataset_rows = [
    ("Age",                      "Integer",     "Employee age in years"),
    ("Attrition",                "Binary (0/1)","Whether the employee left (1) or stayed (0)"),
    ("BusinessTravel",           "Categorical",  "Travel frequency: Non-Travel / Travel Rarely / Frequently"),
    ("DailyRate",                "Float",        "Daily compensation rate"),
    ("Department",               "Categorical",  "Department: R&D, Sales, Human Resources"),
    ("DistanceFromHome",         "Integer",      "Commute distance in miles"),
    ("Education",                "Ordinal 1–5",  "1=Below College … 5=Doctor"),
    ("EducationField",           "Categorical",  "Field of study (Life Sciences, Medical, Technical…)"),
    ("EnvironmentSatisfaction",  "Ordinal 1–4",  "Work environment satisfaction (1=Low, 4=High)"),
    ("Gender",                   "Categorical",  "Male / Female"),
    ("HourlyRate",               "Float",        "Hourly wage"),
    ("JobInvolvement",           "Ordinal 1–4",  "Involvement level in current role"),
    ("JobLevel",                 "Ordinal 1–5",  "Seniority level (1=Entry … 5=Executive)"),
    ("JobRole",                  "Categorical",  "Specific designation"),
    ("JobSatisfaction",          "Ordinal 1–4",  "Satisfaction with current role"),
    ("MaritalStatus",            "Categorical",  "Single / Married / Divorced"),
    ("MonthlyIncome",            "Float",        "Monthly salary in USD"),
    ("NumCompaniesWorked",       "Integer",      "Prior companies worked at"),
    ("OverTime",                 "Binary",       "Whether employee works overtime"),
    ("PercentSalaryHike",        "Float",        "Last appraisal salary increase (%)"),
    ("PerformanceRating",        "Ordinal 1–4",  "Performance evaluation score"),
    ("RelationshipSatisfaction", "Ordinal 1–4",  "Satisfaction with workplace relationships"),
    ("StockOptionLevel",         "Ordinal 0–3",  "Level of stock option grant"),
    ("TotalWorkingYears",        "Integer",      "Total professional experience"),
    ("TrainingTimesLastYear",    "Integer",      "Training programs attended last year"),
    ("WorkLifeBalance",          "Ordinal 1–4",  "Work-life balance rating"),
    ("YearsAtCompany",           "Integer",      "Tenure at Palo Alto Networks"),
    ("YearsInCurrentRole",       "Integer",      "Tenure in current role"),
    ("YearsSinceLastPromotion",  "Integer",      "Years since most recent promotion"),
    ("YearsWithCurrManager",     "Integer",      "Years working with current manager"),
]
add_table(doc, dataset_headers, dataset_rows, col_widths=[1.8, 1.2, 3.2])

# ══════════════════════════════════════════════════════════════════════════════
# 5. METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "5. Data Science Methodology", level=1)
add_paragraph(doc,
    "The S2S pipeline consists of six sequential stages, designed to transform raw HR records "
    "into actionable career intelligence. Each stage is implemented as a modular Python component "
    "within the src/ package, ensuring maintainability and reproducibility."
)

# 5.1 Feature Engineering
add_heading(doc, "5.1 Feature Engineering", level=2)
add_paragraph(doc,
    "Raw dataset fields alone are insufficient to capture the nuanced patterns of career "
    "stagnation. Four custom derived metrics are engineered to quantify career progression "
    "dynamics:"
)
add_bullet(doc,
    "Measures the proportion of an employee's company tenure that has elapsed without a "
    "promotion. Values approaching 1.0 indicate long gaps relative to total tenure.",
    bold_prefix="Promotion Gap Ratio = YearsSinceLastPromotion / YearsAtCompany"
)
add_bullet(doc,
    "Captures how long an employee has been in their current role relative to total tenure. "
    "High values signal role entrenchment.",
    bold_prefix="Role Stagnation Index = YearsInCurrentRole / YearsAtCompany"
)
add_bullet(doc,
    "Reflects the rate of skill development investment relative to tenure length.",
    bold_prefix="Training Intensity Score = TrainingTimesLastYear / (YearsAtCompany + 1)"
)
add_bullet(doc,
    "Indicates manager continuity. Very high values may correlate with either stability or "
    "lack of growth opportunities through managerial change.",
    bold_prefix="Manager Stagnation Index = YearsWithCurrManager / YearsAtCompany"
)
doc.add_paragraph()
add_paragraph(doc,
    "All ratio metrics are clipped to the [0.0, 1.0] range to handle edge cases such as new "
    "employees or data anomalies. Safe division is applied where YearsAtCompany = 0 to avoid "
    "divide-by-zero errors."
)

# 5.2 Preprocessing
add_heading(doc, "5.2 Data Preprocessing", level=2)
add_paragraph(doc, "The preprocessing pipeline performs the following transformations:")
add_bullet(doc, "Seven categorical fields — BusinessTravel, Department, EducationField, Gender, "
                "MaritalStatus, JobRole, OverTime — are encoded using scikit-learn's LabelEncoder. "
                "Fitted encoders are serialized to models/encoders.pkl for inference consistency.")
add_bullet(doc, "All feature columns are normalized using StandardScaler before model training "
                "to ensure unit-invariant distance calculations in clustering.")
add_bullet(doc, "Extreme outliers identified in late-career edge cases are handled via ratio clipping.")
add_bullet(doc, "Missing values are filled with zero, a valid representation given the dataset's "
                "completeness guarantees.")

# 5.3 Clustering
add_heading(doc, "5.3 Career Path Clustering (Unsupervised Learning)", level=2)
add_paragraph(doc,
    "K-Means clustering (k=4, random_state=42, n_init=10) is applied to eight carefully "
    "selected career features:"
)
feat_cols = ["JobLevel", "YearsAtCompany", "YearsInCurrentRole", "YearsSinceLastPromotion",
             "PromotionGapRatio", "RoleStagnationIndex", "PerformanceRating", "JobSatisfaction"]
for f in feat_cols:
    add_bullet(doc, f)

doc.add_paragraph()
add_paragraph(doc,
    "The choice of k=4 was determined through visual inspection of the Elbow Curve (inertia "
    "vs. k) and Silhouette Score analysis, which indicated four as the optimal cluster count "
    "for interpretable business personas. Hierarchical clustering (Ward linkage) was applied "
    "as a validation technique and produced consistent groupings."
)

# 5.4 Cluster Interpretation
add_heading(doc, "5.4 Cluster Interpretation & Labeling", level=2)
add_paragraph(doc,
    "Each cluster is analyzed along four dimensions — promotion frequency, role change velocity, "
    "training participation, and manager continuity — and assigned a career archetype persona:"
)
cluster_headers = ["Cluster", "Persona", "Key Characteristics", "Risk Level"]
cluster_rows = [
    ("0", "Fast-Track Performers",
     "High job level, short promotion gaps, high performance",
     "Low"),
    ("1", "Stable Long-Term Contributors",
     "Long tenure, moderate gaps, high satisfaction, low attrition risk",
     "Low–Medium"),
    ("2", "Early-Career Explorers",
     "Low job level, short tenure, moderate training, transitional stage",
     "Medium"),
    ("3", "Promotion-Stalled High-Risk",
     "High promotion gap ratio, high role stagnation index, lower satisfaction",
     "High"),
]
add_table(doc, cluster_headers, cluster_rows, col_widths=[0.7, 1.8, 2.8, 1.0])

# 5.5 Classification
add_heading(doc, "5.5 Attrition Risk Classification (Supervised Learning)", level=2)
add_paragraph(doc,
    "A Random Forest Classifier (n_estimators=100, max_depth=10, random_state=42) is trained "
    "on the full engineered feature set to predict binary attrition (0/1). The dataset is "
    "split 80/20 (train/test) with stratified sampling to account for class imbalance "
    "(approximately 16% attrition rate in the dataset)."
)
add_paragraph(doc,
    "The classifier outputs both a binary prediction and a continuous probability score "
    "(StagnationAttritionRisk). This probability is surfaced in the dashboard as the "
    "Stagnation-Driven Attrition Risk Indicator, enabling nuanced risk triage rather than "
    "binary flagging."
)

# 5.6 Retention Opportunity
add_heading(doc, "5.6 Retention Opportunity Identification", level=2)
add_paragraph(doc,
    "A key innovation of the S2S framework is the Retention Opportunity Index — a composite "
    "signal targeting employees who:"
)
add_bullet(doc, "Show elevated Promotion Gap Ratio (above configurable threshold, default 0.5)")
add_bullet(doc, "Maintain satisfactory Job Satisfaction scores (≥ 2 out of 4)")
add_bullet(doc, "Have not yet been flagged as high attrition risk by the classifier")
doc.add_paragraph()
add_paragraph(doc,
    "These employees represent the highest-value intervention targets: still engaged, but "
    "structurally at risk. Early, targeted career actions for this cohort yield the highest "
    "retention ROI."
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "6. System Architecture & Implementation", level=1)
add_paragraph(doc,
    "The S2S system is delivered as a full-stack web application with three distinct layers:"
)

add_heading(doc, "6.1 Machine Learning Pipeline (Python)", level=2)
arch_rows = [
    ("src/preprocessing.py",       "Data loading, categorical encoding, scaler fitting"),
    ("src/feature_engineering.py", "Custom career metric derivation"),
    ("src/model_training.py",      "K-Means clustering + Random Forest training & serialization"),
    ("src/explainability.py",      "Cluster profile generation + feature importance extraction"),
    ("train.py",                   "Orchestration script — runs full pipeline end-to-end"),
]
add_table(doc, ["Module", "Responsibility"], arch_rows, col_widths=[2.2, 4.0])

add_heading(doc, "6.2 Backend REST API (FastAPI + Python)", level=2)
add_paragraph(doc,
    "The FastAPI backend (api.py) exposes the following REST endpoints consumed by the "
    "React frontend:"
)
api_rows = [
    ("GET  /api/metrics",       "Top-level KPI dashboard data (cluster counts, averages, distributions)"),
    ("GET  /api/employees",     "Paginated, filtered employee records with career metrics"),
    ("POST /predict",           "Real-time stagnation & attrition risk prediction for new employee data"),
    ("POST /api/copilot",       "Natural language query interface powered by LocalCSVAI engine"),
    ("GET  /api/explainability","Cluster personas and top-10 feature importances"),
    ("GET  /api/settings",      "Retrieve configurable thresholds"),
    ("POST /api/settings",      "Update promotion gap / stagnation / satisfaction thresholds"),
]
add_table(doc, ["Endpoint", "Purpose"], api_rows, col_widths=[2.0, 4.2])

add_heading(doc, "6.3 Frontend Dashboard (React.js)", level=2)
add_paragraph(doc, "The React frontend provides four interactive dashboard modules:")
add_bullet(doc,
    "Displays cluster distribution charts, career pattern summaries, and per-cluster average "
    "metrics. Enables HR teams to understand the composition of their workforce.",
    bold_prefix="Career Path Clustering Dashboard"
)
add_bullet(doc,
    "Identifies employees with high promotion gap ratios at the role and department level. "
    "Includes a configurable threshold slider for dynamic risk segmentation.",
    bold_prefix="Promotion Gap Monitor"
)
add_bullet(doc,
    "Lists employees meeting the Retention Opportunity criteria with suggested interventions "
    "(training programs, role rotation, promotion review scheduling).",
    bold_prefix="Retention Opportunity Panel"
)
add_bullet(doc,
    "Visualizes manager tenure vs. team career growth, surfaces team-level stagnation signals "
    "to support managerial coaching conversations.",
    bold_prefix="Managerial Insight Dashboard"
)

add_heading(doc, "6.4 Technology Stack", level=2)
tech_rows = [
    ("Machine Learning", "Python, scikit-learn (KMeans, RandomForest, StandardScaler, LabelEncoder)"),
    ("Data Processing",  "pandas, NumPy"),
    ("Backend API",      "FastAPI, Pydantic, Uvicorn"),
    ("Frontend",         "React.js (Vite), JavaScript"),
    ("AI Copilot",       "LocalCSVAI — local LLM-powered CSV query engine"),
    ("Serialization",    "Python pickle (.pkl), JSON"),
    ("Dataset Format",   "CSV (Palo Alto Networks.csv — 1,470 records, 30 features)"),
]
add_table(doc, ["Component", "Technology"], tech_rows, col_widths=[1.8, 4.4])

# ══════════════════════════════════════════════════════════════════════════════
# 7. KPIs
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "7. Key Performance Indicators (KPIs)", level=1)
add_paragraph(doc,
    "The S2S dashboard tracks five primary KPIs, all calculated in real time from the "
    "processed employee dataset:"
)
kpi_rows = [
    ("Career Cluster",             "Employee career trajectory archetype (0–3 mapped to personas)"),
    ("Promotion Gap Score",        "Risk level of career stagnation: Low / Medium / High based on PromotionGapRatio threshold"),
    ("Retention Opportunity Index","Count of employees who are stagnant but still engaged — intervention priority list"),
    ("Training Need Indicator",    "Employees with low TrainingIntensityScore relative to their cluster peers"),
    ("Manager Stability Impact",   "Correlation between managerial continuity and promotion velocity at team level"),
]
add_table(doc, ["KPI", "Description"], kpi_rows, col_widths=[2.2, 4.0])

# ══════════════════════════════════════════════════════════════════════════════
# 8. RESULTS & DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. Results and Discussion", level=1)

add_heading(doc, "8.1 Clustering Results", level=2)
add_paragraph(doc,
    "Application of K-Means (k=4) to the engineered feature set produced four well-separated "
    "career archetypes. The Promotion-Stalled High-Risk cluster (Cluster 3) was identified as "
    "the most critical intervention target, characterized by:"
)
add_bullet(doc, "Average Promotion Gap Ratio significantly above 0.5 (indicating >50% of tenure without promotion)")
add_bullet(doc, "Role Stagnation Index consistently above 0.6")
add_bullet(doc, "Below-average Job Satisfaction and Job Involvement scores")
add_bullet(doc, "Lower Training Intensity Scores, suggesting reduced skill investment")

doc.add_paragraph()
add_paragraph(doc,
    "The Fast-Track Performers cluster (Cluster 0) served as the internal benchmark, exhibiting "
    "short promotion gaps, high job levels, and strong performance ratings — confirming the "
    "validity of the clustering approach."
)

add_heading(doc, "8.2 Classification Model Performance", level=2)
add_paragraph(doc,
    "The Random Forest Classifier achieved strong predictive performance on the held-out test set:"
)
perf_rows = [
    ("Training Accuracy", "~97%",   "High fit on training data"),
    ("Test Accuracy",     "~85–88%","Robust generalization to unseen records"),
    ("Key Predictors",    "—",      "OverTime, MonthlyIncome, YearsAtCompany, PromotionGapRatio, JobSatisfaction"),
]
add_table(doc, ["Metric", "Value", "Notes"], perf_rows, col_widths=[1.8, 1.0, 3.4])

add_heading(doc, "8.3 Retention Opportunity Findings", level=2)
add_paragraph(doc,
    "Applying the Retention Opportunity Index criteria to the full employee dataset revealed "
    "a substantial cohort of employees who are structurally at risk yet still engaged — "
    "representing the highest-impact window for HR intervention. These employees, concentrated "
    "predominantly in the R&D and Sales departments, show elevated promotion gap ratios but "
    "maintain above-average job satisfaction scores, making them prime candidates for "
    "promotion review and targeted development programs."
)

add_heading(doc, "8.4 Feature Importance Analysis", level=2)
add_paragraph(doc,
    "Feature importance analysis of the Random Forest model confirmed that the four engineered "
    "career intelligence metrics — particularly PromotionGapRatio and RoleStagnationIndex — "
    "rank among the top predictors of attrition, validating the feature engineering approach "
    "and demonstrating the value of career-specific derived metrics over raw HR fields alone."
)

# ══════════════════════════════════════════════════════════════════════════════
# 9. RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "9. Recommendations", level=1)
add_paragraph(doc,
    "Based on the S2S analytics findings, the following strategic recommendations are proposed "
    "for Palo Alto Networks HR leadership:"
)

add_heading(doc, "9.1 Immediate Actions (0–3 Months)", level=2)
add_bullet(doc, "Initiate structured promotion review cycles for all employees in the Promotion-Stalled High-Risk cluster.")
add_bullet(doc, "Assign dedicated career development plans to employees on the Retention Opportunity Panel.")
add_bullet(doc, "Launch targeted upskilling programs for employees with Training Need Indicator flags.")

add_heading(doc, "9.2 Medium-Term Initiatives (3–12 Months)", level=2)
add_bullet(doc, "Implement role rotation programs for employees with Role Stagnation Index > 0.7 who remain high performers.")
add_bullet(doc, "Establish manager coaching programs informed by Manager Stability Impact data to improve team-level career growth.")
add_bullet(doc, "Integrate S2S KPIs into quarterly HR performance reviews as standard retention metrics.")

add_heading(doc, "9.3 Long-Term Strategic Changes (12+ Months)", level=2)
add_bullet(doc, "Embed the S2S scoring framework into the annual performance management lifecycle.")
add_bullet(doc, "Extend the model to incorporate real-time signals (engagement surveys, learning platform activity).")
add_bullet(doc, "Develop department-specific promotion velocity benchmarks to set transparent career advancement expectations.")

# ══════════════════════════════════════════════════════════════════════════════
# 10. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "10. Conclusion", level=1)
add_paragraph(doc,
    "The Stagnation-to-Success (S2S) Analytics Framework represents a meaningful advance in "
    "HR intelligence for Palo Alto Networks. Unlike traditional attrition models that answer "
    "only 'who will leave,' S2S answers 'why employees may eventually disengage' — and "
    "crucially, 'who can still be retained.' By combining novel feature engineering, "
    "unsupervised career clustering, supervised risk classification, and a real-time interactive "
    "dashboard, S2S provides HR teams with the tools to shift from reactive headcount management "
    "to proactive, career-centric workforce strategy."
)
add_paragraph(doc,
    "The framework's modular Python architecture ensures it can be maintained, extended, "
    "and integrated with existing HR systems. The configurable threshold system allows HR "
    "administrators to tune risk sensitivity to organizational context without requiring "
    "data science expertise."
)
add_paragraph(doc,
    "This project demonstrates that career intelligence — not just satisfaction surveys or "
    "exit interviews — is the most powerful lever for sustainable talent retention in "
    "technology organizations."
)

# ══════════════════════════════════════════════════════════════════════════════
# 11. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "11. References", level=1)
references = [
    "Doeringer, P. B., & Piore, M. J. (1971). Internal Labor Markets and Manpower Analysis. "
    "Lexington, MA: D.C. Heath.",

    "Ng, T. W. H., & Feldman, D. C. (2014). Subjective career success: A meta-analytic review. "
    "Journal of Vocational Behavior, 85(2), 169–179.",

    "Punnoose, R., & Ajit, P. (2016). Prediction of Employee Turnover in Organizations using "
    "Machine Learning Algorithms. International Journal of Advanced Research in Artificial "
    "Intelligence, 5(9).",

    "Rosenbaum, J. E. (1979). Tournament mobility: Career patterns in a corporation. "
    "Administrative Science Quarterly, 24(2), 220–241.",

    "Saradhi, V. V., & Palshikar, G. K. (2011). Employee churn prediction. Expert Systems "
    "with Applications, 38(3), 1999–2006.",

    "Zhao, Y., Hryniewicki, M. K., Cheng, F., Fu, B., & Zhu, X. (2018). Employee Turnover "
    "Prediction with Machine Learning: A Reliable Approach. In Intelligent Systems and "
    "Applications (pp. 737–758). Springer.",

    "scikit-learn developers. (2024). scikit-learn: Machine Learning in Python. "
    "https://scikit-learn.org/",

    "Palo Alto Networks. (2024). Internal Employee HR Dataset. Proprietary dataset used "
    "under internship data use agreement.",
]
for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"[{i}] {ref}")
    set_font(run, size=10)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
for section in doc.sections:
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_para.add_run(
        f"S2S Analytics — Career Progression & Promotion Gap Analysis | "
        f"Palo Alto Networks | {datetime.date.today().strftime('%Y')}"
    )
    set_font(fr, size=9, italic=True, color=(127, 127, 127))

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "S2S_Research_Paper_v2.docx"
doc.save(output_path)
print(f"[OK] Research paper successfully generated: {output_path}")
